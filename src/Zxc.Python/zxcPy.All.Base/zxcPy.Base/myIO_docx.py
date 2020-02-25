# -*- coding: utf-8 -*-
"""
Created on  张斌 2020-01-09 10:00:00 
    @author: zhang bin
    @email:  zhangbin@gsafety.com

    Word操作 
    @依赖库： python-docx
"""
import sys, os, datetime, codecs
from docx import Document
from docx.shared import Inches

#加载自定义库
import myEnum, myData, myIO, myData_Trans    

#定义数据结构枚举
myDcox_paragraph_Type = myEnum.enum('Title', 'Heading', 'Intense_Quote', 'List_Bullet', 'List_Number', 'Normal', 'Other')


# Docx段落信息
class myDcox_paragraph:
    def __init__(self, strTitle = ""):
        self.type = myDcox_paragraph_Type.Other
        self.level = 0          # 段落级别
        self.valid = True       # 段落有效，区分虚拟段落
        #self.headID = -1        # 段落段落的标题 ID 
        #self.parentID = -1      # 父级ID
        self.headTitle = ""     # 标题名称
        self.headID = 0         # 标题ID
        self.headID_str = ""    # 标题ID（父级级联或设定）
        self.headCount_child = 0# 标题子级别总数
        self.contents = []      # 内容
        self.Childs = []        # 子集
        self.Parent = None      # 父级对象
        self.Link_str = ""      # 返回链接信息
        self.Out_Title = True   # 允许输出标题信息
        if(strTitle != ""): self.setTitle(strTitle)
        
    # 设置标题信息
    def setTitle(self, strTitle):
        pass
    
    # 新增子节点
    def addChild(self, child):
        if(child == None): return False
        self.Childs.append(child)
        child.Parent = self
        self.updatAll(self.headID)
        return True
    
    # 子集更新
    def updatAll(self, nId = 0):
        self.updatLevel()

        if(nId < 1): nId = self.level
        self.updatID(nId)
    # 更新级别
    def updatLevel(self, bCheck = False):
        if(bCheck == False): return

        if(self.Parent == None):
            self.level = 1
        else:
            if(self.Parent != None):
                self.level = self.Parent.level + 1

        # 更新子集
        for x in self.Childs:
            x.updatLevel()
    def updatID(self, nID = 1):
        self.headID = myData.iif(nID < 1, 1, nID)

        # 修改子集
        ind = 1
        for x in self.Childs:
            x.updatID(ind)
            ind += 1

    # 提取ID索引号
    def getheadID_str(self):
        self.headID_str = myData.iif(self.headID < 1, "", str(self.headID))

        # 组装父项ID
        strID_pre = ""
        parent = self.Parent 
        while(parent != None):
            strID_pre = str(parent.headID) + "." + strID_pre
            parent = parent.Parent
        self.headID_str = strID_pre + self.headID_str
        return self.headID_str

# Docx文档信息
class myDcox:
    def __init__(self, path=""):
        self.parasDocx = []                 #段落信息集（不含主标题，不含索引）
        self.paraDef = myDcox_paragraph()   #添加默认项避免无节点报错 
        self.paraCurrent = self.paraDef 
        self.document = None
        self.paraDepth_min = 1              # 段落级别深度-最小
        self.paraDepth_max = 0              # 段落级别深度-最大
        self.Ind = 1 
        if(path != ""): self.loadDocx(path)
        
    # 加载Docx文件 
    def loadDocx(self, path, onlyHeading = False):
        # 文件必须存在 
        if (os.path.exists(path) == False): return False
        
        # 打开文件docx，解析Docx段落集 
        self.document = Document(path)  
        return self._loadDocx_paras(self.document, onlyHeading)
    # 加载Docx段落集 
    def _loadDocx_paras(self, document, onlyHeading = False):
        # 提取所有行数据、解析所有
        ind = 0
        nNums = len(document.paragraphs)
        for i in range(0, nNums):
            #if i < ind: continue
            paragraph = document.paragraphs[i]
            para = myDcox_paragraph()

            # 区分段落类型
            paraStyle =  paragraph.style.name.replace(' ', '_')
            if(paraStyle == myDcox_paragraph_Type.Title):
                #主标题、题目
                para.type = myDcox_paragraph_Type.Title
                self.paraDef = para
                pass
            elif(paraStyle.count(myDcox_paragraph_Type.Heading) == 1):
                #各级标题
                para.type = myDcox_paragraph_Type.Heading
                para.level = int(paraStyle.replace(myDcox_paragraph_Type.Heading + "_", ""))
                para.headTitle = paragraph.text
                self.addParagraph(para)
                pass
            else:
                if(onlyHeading): continue

                # 其他类型处置，未完善
                para.type = paraStyle
                self.paraCurrent.contents.append(paragraph.text)
            
    # 添加段落
    def addParagraph(self, para):
        if(self.updataCurrent_para(para) == 0):
            self.parasDocx.append(para)
        
        # 更新当前段落级别
        para.updatLevel()
        if(self.paraDepth_max< para.level):
            self.paraDepth_max = para.level
        
    # 更新当前节点  
    def updataCurrent_para(self, para):
        # 初始记录节点信息
        if(self.paraCurrent == None or para.level <= 1):
            self.paraCurrent = para
            return 0

        # 段落无顶级的情况，补齐节点到第一级
        if(self.paraCurrent.level == 0 and para.level > 1):
            paraTemp = para
            while(paraTemp.level > 1):
                self.paraCurrent = myDcox_paragraph()
                self.paraCurrent.level = paraTemp.level - 1
                self.paraCurrent.valid = False                  #虚拟段落
                self.paraCurrent.addChild(paraTemp)

                # 第一级退出
                if(self.paraCurrent.level == 1):
                    self.parasDocx.append(self.paraCurrent)
                    break
                paraTemp = self.paraCurrent
            self.paraCurrent = para 
            self.paraDepth_min = para.level
            return 1
        
        # 更新节点信息
        if(self.paraCurrent.level < para.level):
            # 当前标题级别比前一个小，则为子ID
            self.paraCurrent.addChild(para)
            self.paraCurrent = para
        else:
            # 同级或更高级(非顶级)
            parent = self.paraCurrent.Parent
            while(parent != None and parent.level >= para.level):
                parent = parent.Parent

            # 更新
            if(parent != None):
                parent.addChild(para)
                self.paraCurrent = para
        return 1
    
    # 提取所有段落信息集合
    def getParas_All(self, paraCur = None, onlyValid = True):
        # 提取子集
        if(paraCur == None): 
            paras= pDocx.parasDocx
        else: 
            paras = paraCur.Childs
            
        lstParas = []
        for para in paras:
            if(onlyValid == False or para.valid):
                lstParas.append(para)

            # 递归子集
            lstParas_temp = self.getParas_All(para, onlyValid)
            lstParas += lstParas_temp
            para.headCount_child = self.getParas_count(lstParas_temp)
        if(paraCur != None):
            paraCur.headCount_child = self.getParas_count(lstParas)
        return lstParas
    
    # 提取底层段落数
    def getParas_count(self, lstParas):
        nCount = 0
        for x in lstParas:
            if(x.headCount_child == 0):
                nCount += 1
        if(nCount > 0):
             aa= 0
        return nCount

# 设置标题信息
def checkLine(strLine, strTag=""):
    document = Document()

    #插入标题
    document.add_heading('Document Title', 0)  

    # 插入段落,单字特殊风格
    p = document.add_paragraph('A plain paragraph having some ')   
    p.add_run('bold').bold = True
    p.add_run(' and some ')
    p.add_run('italic.').italic = True

    # 插入标题,插入内容（风格）
    document.add_heading('Heading, level 1', level=1)
    document.add_paragraph('Intense quote', style='IntenseQuote')

    # 插入列表（符号列表、编号列表）
    document.add_paragraph(
        'first item in unordered list', style='ListBullet'
    )
    document.add_paragraph(
        'first item in ordered list', style='ListNumber'
    )

    # 插入图片
    # document.add_picture('monty-truth.png', width=Inches(1.25)) 

    # 插入表格
    table = document.add_table(rows=1, cols=3) 
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Qty'
    hdr_cells[1].text = 'Id'
    hdr_cells[2].text = 'Desc'
    for item in ['1', '2']:
        row_cells = table.add_row().cells
        row_cells[0].text = str(item) + "-列1"
        row_cells[1].text = str(item) + "-列2"
        row_cells[2].text = str(item) + "-列desc"

    # 插入分页符
    document.add_page_break()

    # 保存文档
    document.save('data/demo.docx')  

    #读取表格材料，并输出结果
    tables = [table for table in document.tables];
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                print( cell.text,'\t')
            print
        print ('\n')


if __name__ == '__main__': 
    filename= "data/应急救援.docx"
    #pDocx = myDcox(filename)
    pDocx = myDcox()
    pDocx.loadDocx(filename, True)

    # 提取测试 

    # 导出节点为excel
    paras = pDocx.getParas_All(None, True)
    if(len(paras) > 0):
        import myIO_xlsx
        pExcel = myIO_xlsx.DtTable()
        pExcel.dataName = "应急救援"
        for x in range(pDocx.paraDepth_min, pDocx.paraDepth_max + 1):
            pExcel.dataField.append("段落 " + str(x) + " 级")
        pExcel.dataField.append("备注")
    
        styleSets = []
        nDelta = pDocx.paraDepth_min
        nCols = pExcel._Cloumns()
        ind = 1             # 第一行不算
        nRows = 0
        for x in range(0, len(paras)):
            #循环跳过，
            if(nRows > 0):
                nRows -= 1
                continue

            #子行合并
            para = paras[x]
            values = [""] * nCols

            while(True):
                para = paras[x + nRows]
                indCol = para.level - nDelta
                values[indCol] = para.headTitle
                values[nCols - 1] = "《应急救援.docx》段落 " + para.getheadID_str()
                
                # 主节点合并
                if(para.headCount_child > 0):
                    styleSet = {'type': "Merge", 'pos': [ind, ind + para.headCount_child - 1, indCol, indCol], 'text': para.headTitle}
                    styleSets.append(styleSet)

                # 退出于跟节点
                if(para.headCount_child == 0):
                    break
                nRows += 1

            # 加入到表格
            pExcel.dataMat.append(values)
            ind += 1

        # 保存
        pExcel.Save("D:\\myCode\\zxcProj\\src\\Zxc.Python\\zxcPy.All.Base\\Data", "应急救援功能点", styleSets=styleSets)

    print()


